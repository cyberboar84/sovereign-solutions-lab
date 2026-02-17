#!/usr/bin/env python3
"""Convert all Test 5 translation files to .docx format."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = "."

FILES = [
    ("essay_es_v4.md", "essay_spanish.docx", "Metáforas en El Gran Gatsby: Un Análisis Literario"),
    ("essay_zh_v4.md", "essay_mandarin.docx", "了不起的盖茨比中的隐喻：文学分析"),
    ("summary_es_v4.md", "summary_spanish.docx", "Resumen Ejecutivo"),
    ("summary_zh_v4.md", "summary_mandarin.docx", "执行摘要"),
]

for src_name, dst_name, title in FILES:
    src = os.path.join(BASE, src_name)
    dst = os.path.join(BASE, dst_name)

    if not os.path.exists(src):
        print(f"SKIP: {src_name} not found")
        continue

    with open(src, "r", encoding="utf-8") as f:
        text = f.read()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process markdown
    for section in text.split("## "):
        if not section.strip():
            continue

        lines = section.strip().split("\n")
        heading_text = lines[0].strip().replace("# ", "").replace("*", "")

        # Skip if this is the main title (already added)
        if heading_text.startswith("Metáforas") or heading_text.startswith("了不起"):
            continue

        # Add section heading
        if not heading_text.startswith("#"):
            doc.add_heading(heading_text, level=2)

        # Add body paragraphs
        body = "\n".join(lines[1:]).strip()
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    doc.save(dst)
    print(f"Saved: {dst_name} ({os.path.getsize(dst) // 1024} KB)")

print("\nAll translations converted to .docx")
print("Download via: python3 -m http.server 8888")
